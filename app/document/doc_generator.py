"""
doc_generator.py
Takes the final generated sections and produces a polished .docx file
using python-docx: title page styling, headings, a horizontal rule effect,
and consistent professional formatting.
"""

import os
import uuid
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.config import settings

ACCENT_COLOR = RGBColor(0x1E, 0x3A, 0x5F)   # deep navy
SUBTLE_COLOR = RGBColor(0x5A, 0x5A, 0x5A)   # gray


def _add_bottom_border(paragraph):
    """Adds a thin bottom border under a paragraph (used under the title)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1E3A5F")
    p_borders.append(bottom)
    p_pr.append(p_borders)


def build_document(document_title: str, document_type: str, assumptions: list,
                    sections: dict, user_request: str) -> str:
    """
    Builds the .docx file and saves it to the outputs directory.
    Returns the filename (not full path) of the generated document.
    """
    doc = Document()

    # --- Base style tweaks ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # --- Title block ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(document_title)
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = ACCENT_COLOR

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run(document_type.replace("_", " ").title())
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.color.rgb = SUBTLE_COLOR
    _add_bottom_border(subtitle_para)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(f"Generated on {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = SUBTLE_COLOR

    doc.add_paragraph()  # spacing

    # --- Original request context ---
    doc.add_heading("Request Summary", level=2)
    p = doc.add_paragraph(user_request)
    p.runs[0].font.color.rgb = SUBTLE_COLOR

    # --- Assumptions (only if the agent had to make any) ---
    if assumptions:
        doc.add_heading("Assumptions Made by the Agent", level=2)
        for a in assumptions:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_paragraph()

    # --- Main generated sections ---
    for section_title, content in sections.items():
        doc.add_heading(section_title, level=1)
        for para_text in content.split("\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            if para_text.startswith(("- ", "* ")):
                doc.add_paragraph(para_text[2:], style="List Bullet")
            else:
                doc.add_paragraph(para_text)

    # --- Footer note ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "Generated automatically by an autonomous AI agent (FastAPI + Gemini)."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = SUBTLE_COLOR

    # --- Save ---
    filename = f"{document_type}_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.join(settings.OUTPUT_DIR, filename)
    doc.save(filepath)

    return filename
